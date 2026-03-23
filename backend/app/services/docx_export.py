"""
Converts a Markdown string to a .docx file and returns a BytesIO object.

Supported Markdown syntax:
  - # Heading 1 / ## Heading 2 / ### Heading 3
  - **bold** and *italic* (inline, within paragraphs)
  - - item  or  * item  (unordered lists)
  - 1. item  (ordered lists)
  - Plain paragraphs
"""

import re
import io
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_run_with_formatting(paragraph, text: str) -> None:
    """
    Parses inline **bold** and *italic* markers and adds appropriately
    formatted runs to *paragraph*.
    """
    # Pattern captures bold (**...**), italic (*...*), or plain text
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)')
    for match in pattern.finditer(text):
        bold_text, italic_text, plain_text = match.groups()
        if bold_text is not None:
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif italic_text is not None:
            run = paragraph.add_run(italic_text)
            run.italic = True
        else:
            paragraph.add_run(plain_text)


def markdown_to_docx(markdown: str) -> io.BytesIO:
    """
    Convert *markdown* text to a docx document and return it as a BytesIO.
    """
    doc = Document()

    # Set default body font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = markdown.splitlines()
    i = 0
    list_counter = 0  # tracks current number for ordered lists

    while i < len(lines):
        line = lines[i]

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            list_counter = 0

        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            list_counter = 0

        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            list_counter = 0

        # ── Unordered list item ───────────────────────────────────────────────
        elif re.match(r'^[-*]\s+', line):
            content = re.sub(r'^[-*]\s+', '', line)
            para = doc.add_paragraph(style='List Bullet')
            _add_run_with_formatting(para, content)
            list_counter = 0

        # ── Ordered list item ─────────────────────────────────────────────────
        elif re.match(r'^\d+\.\s+', line):
            content = re.sub(r'^\d+\.\s+', '', line)
            para = doc.add_paragraph(style='List Number')
            _add_run_with_formatting(para, content)
            list_counter += 1

        # ── Blank line ────────────────────────────────────────────────────────
        elif line.strip() == '':
            list_counter = 0

        # ── Regular paragraph ─────────────────────────────────────────────────
        else:
            para = doc.add_paragraph()
            _add_run_with_formatting(para, line)
            list_counter = 0

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
